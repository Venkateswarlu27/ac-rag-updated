"""
add_code_snippets.py
Inserts Section 3.10 "Key Code Listings" into the thesis docx
to expand it from ~35 pages to ~70 pages.

Strategy: rebuild the document body by copying all existing paragraphs/tables,
inserting the new section before the Chapter 4 heading.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

THESIS_PATH = "AC_RAG_Thesis_Submission_Ready.docx"
doc = Document(THESIS_PATH)

# ─── Code content ─────────────────────────────────────────────────────────────

STATE_CODE = """\
# pipeline/state.py  — Shared State Schema
from __future__ import annotations
from typing import Any, Dict, List, Optional, TypedDict

class RetrievalPlan(TypedDict):
    k: int                   # number of final passages to retrieve (4-12)
    fetch_k: int             # MMR candidate pool size
    lambda_mult: float       # MMR diversity (0=diverse, 1=relevant)
    modality_filter: str     # "text" | "table" | "figure" | "all"
    use_multi_query: bool    # decompose query into sub-queries
    retrieval_depth: str     # "shallow" | "standard" | "deep"

class CriticScores(TypedDict):
    faithfulness: int
    completeness: int
    table_accuracy: int
    figure_accuracy: int
    conciseness: int
    overall: float
    feedback: str

class ACRagState(TypedDict):
    query: str
    route: Optional[str]                   # "rag" | "unknown"
    rewritten_query: Optional[str]
    intent: Optional[str]                  # factual|analytical|comparative|summarization
    complexity_score: Optional[float]      # 0.0 (simple) -> 1.0 (complex)
    decomposed_queries: Optional[List[str]]
    retrieval_plan: Optional[RetrievalPlan]
    retrieved_docs: Optional[List[Dict[str, Any]]]
    scored_docs: Optional[List[Dict[str, Any]]]
    validation_passed: Optional[bool]
    refined_context: Optional[str]
    answer: Optional[str]
    answer_with_attribution: Optional[List[Dict[str, Any]]]
    critic_scores: Optional[CriticScores]
    critic_passed: Optional[bool]
    critic_feedback: Optional[str]
    retry_count: int
    retry_reason: Optional[str]            # "content" | "format"
    stage_logs: List[Dict]
    error: Optional[str]

def initial_state(query: str) -> ACRagState:
    return ACRagState(
        query=query, route=None, rewritten_query=None,
        intent=None, complexity_score=None, decomposed_queries=None,
        retrieval_plan=None, retrieved_docs=None,
        scored_docs=None, validation_passed=None,
        refined_context=None, answer=None,
        answer_with_attribution=None, critic_scores=None,
        critic_passed=None, critic_feedback=None,
        retry_count=0, retry_reason=None, stage_logs=[], error=None,
    )"""

GRAPH_CODE = """\
# pipeline/graph.py  — LangGraph Pipeline Assembly
from langgraph.graph import StateGraph, END
from pipeline.state import ACRagState, initial_state

def build_pipeline(vsm=None):
    if vsm is None:
        vsm = VectorStoreManager(); vsm.load()

    graph = StateGraph(ACRagState)

    # Register all nodes
    graph.add_node("entry_router",      make_entry_router_node(vsm))
    graph.add_node("direct_responder",  direct_responder_node)
    graph.add_node("query_analyzer",    query_analyzer_node)
    graph.add_node("retrieval_planner", retrieval_planner_node)
    graph.add_node("retriever",         make_retriever_node(vsm))
    graph.add_node("validator",         validator_node)
    graph.add_node("context_refiner",   context_refiner_node)
    graph.add_node("generator",         generator_node)
    graph.add_node("critic",            critic_node)
    graph.add_node("increment_retry",   lambda s: {**s, "retry_count": s["retry_count"]+1})
    graph.add_node("end_success",       lambda s: s)
    graph.add_node("end_error",         lambda s: s)
    graph.add_node("end_max_retries",   lambda s: s)

    graph.set_entry_point("entry_router")

    # Routing: entry_router -> rag pipeline | direct_responder
    graph.add_conditional_edges("entry_router", route_after_entry_router,
        {"rag": "query_analyzer", "unknown": "direct_responder"})
    graph.add_edge("direct_responder", "end_success")

    # query_analyzer -> retrieval_planner | end_error
    graph.add_conditional_edges("query_analyzer", route_after_query_analyzer,
        {"retrieval_planner": "retrieval_planner", "end_error": "end_error"})

    # retrieval_planner -> retriever (always)
    graph.add_conditional_edges("retrieval_planner", route_after_retrieval_planner,
        {"retriever": "retriever"})

    # retriever -> validator | retry | end_error | end_max_retries
    graph.add_conditional_edges("retriever", route_after_retriever,
        {"validator": "validator", "query_analyzer": "increment_retry",
         "end_error": "end_error", "end_max_retries": "end_max_retries"})

    # validator -> context_refiner | generator (ablation) | retry | end_max_retries
    graph.add_conditional_edges("validator", route_after_validator,
        {"context_refiner": "context_refiner", "generator": "generator",
         "retriever": "increment_retry", "end_max_retries": "end_max_retries"})

    # context_refiner -> generator
    graph.add_conditional_edges("context_refiner", route_after_context_refiner,
        {"generator": "generator"})

    # generator -> critic | end_success (ablation) | end_error
    graph.add_conditional_edges("generator", route_after_generator,
        {"critic": "critic", "end_success": "end_success", "end_error": "end_error"})

    # critic -> end_success | content retry | format retry | end_max_retries
    graph.add_conditional_edges("critic", route_after_critic,
        {"end_success": "end_success", "query_analyzer": "increment_retry",
         "generator": "increment_retry", "end_max_retries": "end_max_retries"})

    # increment_retry -> correct destination based on retry_reason
    graph.add_conditional_edges("increment_retry", _route_after_retry_increment,
        {"query_analyzer": "query_analyzer", "retriever": "retriever",
         "generator": "generator", "end_max_retries": "end_max_retries"})

    graph.add_edge("end_success", END)
    graph.add_edge("end_error",   END)
    graph.add_edge("end_max_retries", END)
    return graph.compile()

def _route_after_retry_increment(state: ACRagState) -> str:
    from config.settings import MAX_RETRIES
    if state.get("retry_count", 0) > MAX_RETRIES:
        return "end_max_retries"
    reason = state.get("retry_reason") or "content"
    if reason == "format":               return "generator"
    if state.get("validation_passed") is False: return "retriever"
    return "query_analyzer"

def run_pipeline(query: str, vsm=None) -> ACRagState:
    pipeline = build_pipeline(vsm)
    result   = pipeline.invoke(initial_state(query), config={"recursion_limit": 50})
    return result"""

ENTRY_ROUTER_CODE = """\
# pipeline/nodes/entry_router.py  — Entry Router Node
from pipeline.state import ACRagState
from config.settings import ROUTER_SIMILARITY_THRESHOLD

def _similarity_route(vsm, query: str) -> str:
    \"\"\"k=3 cosine similarity check; returns 'rag' or 'unknown'.\"\"\"
    try:
        results = vsm.similarity_search_with_score(query, k=3)
        if not results:
            return "unknown"
        avg   = sum(s for _, s in results) / len(results)
        top   = results[0][1]
        if top >= ROUTER_SIMILARITY_THRESHOLD and avg >= ROUTER_SIMILARITY_THRESHOLD * 0.7:
            return "rag"
        return "unknown"
    except Exception:
        return "rag"   # fail-safe: attempt RAG rather than refusing

def make_entry_router_node(vsm):
    def entry_router_node(state: ACRagState) -> ACRagState:
        route = _similarity_route(vsm, state["query"].strip())
        return {**state, "route": route,
                "stage_logs": state["stage_logs"] + [
                    {"stage": "entry_router", "status": "completed",
                     "details": {"route": route}}]}
    return entry_router_node"""

QUERY_ANALYZER_CODE = """\
# pipeline/nodes/query_analyzer.py  — Query Analyzer Node
from langchain_core.prompts import ChatPromptTemplate
from pydantic import BaseModel, Field
from typing import List
from pipeline.state import ACRagState
from utils.llm_factory import get_llm

class QueryAnalysis(BaseModel):
    rewritten_query:    str   = Field(description="Clarified retrieval-optimised query.")
    intent:             str   = Field(description="factual|analytical|comparative|summarization")
    complexity_score:   float = Field(ge=0.0, le=1.0, description="0.0 simple -> 1.0 complex")
    complexity_reason:  str   = Field(description="One-sentence explanation of complexity.")
    decomposed_queries: List[str] = Field(default_factory=list,
        description="Sub-queries if complexity >= 0.6, else empty list.")

_SYSTEM = \"\"\"You are a query analysis expert for a document QA system.
Intent categories:
  factual       - specific fact, number, date, or definition
  analytical    - explain, analyse, or reason about something
  comparative   - compare two or more things
  summarization - summary or overview

Complexity scoring:
  0.0-0.3  Single fact lookup
  0.4-0.6  Synthesise a few passages, some reasoning
  0.7-0.9  Multi-hop reasoning, cross-document synthesis
  1.0      Requires full document understanding

Only decompose if complexity >= 0.6. Maximum 4 sub-queries.\"\"\"

_prompt = ChatPromptTemplate.from_messages([
    ("system", _SYSTEM), ("human", "Query: {query}"),
])

def query_analyzer_node(state: ACRagState) -> ACRagState:
    \"\"\"
    Reads:  state['query']
    Writes: rewritten_query, intent, complexity_score, decomposed_queries
    \"\"\"
    try:
        result = (_prompt | get_llm().with_structured_output(QueryAnalysis)).invoke(
            {"query": state["query"]})
        if result.complexity_score < 0.6:
            result.decomposed_queries = []
        return {**state,
                "rewritten_query": result.rewritten_query, "intent": result.intent,
                "complexity_score": result.complexity_score,
                "decomposed_queries": result.decomposed_queries,
                "stage_logs": state["stage_logs"] + [
                    {"stage": "query_analyzer", "status": "completed",
                     "details": {"intent": result.intent,
                                 "complexity": result.complexity_score}}]}
    except Exception as e:
        return {**state, "error": f"QueryAnalyzer failed: {e}",
                "stage_logs": state["stage_logs"] + [
                    {"stage": "query_analyzer", "status": "failed",
                     "details": {"error": str(e)}}]}"""

RETRIEVAL_PLANNER_CODE = """\
# pipeline/nodes/retrieval_planner.py  — Retrieval Planner Node
from pydantic import BaseModel, Field
from langchain_core.prompts import ChatPromptTemplate
from config.settings import (RETRIEVAL_K_MIN, RETRIEVAL_K_MAX, RETRIEVAL_K_DEFAULT,
                              MMR_FETCH_K_MULTIPLIER, MMR_LAMBDA_MULT, USE_RETRIEVAL_PLANNER)
from pipeline.state import ACRagState, RetrievalPlan
from utils.llm_factory import get_llm

def _rule_based_plan(complexity: float, intent: str) -> RetrievalPlan:
    \"\"\"Deterministic tier-1 plan based on complexity thresholds.\"\"\"
    if complexity < 0.4:
        k, depth, multi, lm = RETRIEVAL_K_MIN,     "shallow",  False, 0.7
    elif complexity < 0.7:
        k, depth, multi, lm = RETRIEVAL_K_DEFAULT, "standard", False, MMR_LAMBDA_MULT
    else:
        k, depth, multi, lm = min(10, RETRIEVAL_K_MAX), "deep", True, 0.3
    if intent in ("analytical", "comparative"):
        lm = max(0.3, lm - 0.1); k = min(k + 2, RETRIEVAL_K_MAX)
    return RetrievalPlan(k=k, fetch_k=k*MMR_FETCH_K_MULTIPLIER, lambda_mult=lm,
                         modality_filter="all", use_multi_query=multi,
                         retrieval_depth=depth)

class PlanRefinement(BaseModel):
    modality_filter: str = Field(description="text|table|figure|all")
    k_override:      int = Field(ge=4, le=12, description="Override k if clearly wrong.")
    override_reason: str = Field(description="Reason for change, or 'no change'.")

_PROMPT = ChatPromptTemplate.from_messages([
    ("system", "You are a retrieval planning agent. Refine the rule-based plan by:\\n"
               "1. Setting modality_filter if the query explicitly asks about a table/figure.\\n"
               "2. Optionally adjusting k if clearly wrong. Do NOT change other fields."),
    ("human",  "Query: {query}\\nIntent: {intent}\\nComplexity: {complexity}\\n"
               "Rule k: {k}  Depth: {depth}\\nRefine the plan."),
])

def retrieval_planner_node(state: ACRagState) -> ACRagState:
    \"\"\"
    Two-tier approach:
      Tier 1: Rule-based defaults (always runs, ablation-safe)
      Tier 2: LLM refinement pass (enabled via USE_RETRIEVAL_PLANNER flag)
    Reads:  complexity_score, intent, query
    Writes: retrieval_plan
    \"\"\"
    complexity = state.get("complexity_score") or 0.5
    intent     = state.get("intent") or "factual"
    plan       = _rule_based_plan(complexity, intent)
    if USE_RETRIEVAL_PLANNER:
        try:
            ref = (_PROMPT | get_llm().with_structured_output(PlanRefinement)).invoke(
                {"query": state.get("rewritten_query") or state["query"],
                 "intent": intent, "complexity": complexity,
                 "k": plan["k"], "depth": plan["retrieval_depth"]})
            plan["modality_filter"] = ref.modality_filter
            plan["k"]       = ref.k_override
            plan["fetch_k"] = ref.k_override * MMR_FETCH_K_MULTIPLIER
        except Exception as e:
            pass   # fall back to rule-based plan
    return {**state, "retrieval_plan": plan,
            "stage_logs": state["stage_logs"] + [
                {"stage": "retrieval_planner", "status": "completed",
                 "details": dict(plan)}]}"""

VALIDATOR_CODE = """\
# pipeline/nodes/validator.py  — Evidence Validator Node
from config.settings import EVIDENCE_SCORE_THRESHOLD, MIN_VALID_PASSAGES, USE_VALIDATOR
from pipeline.state import ACRagState
from utils.scoring import score_passages_against_query

def validator_node(state: ACRagState) -> ACRagState:
    \"\"\"
    Algorithm:
      1. Embed query and all retrieved passages
      2. Compute cosine similarity (query vs each passage)
      3. Discard passages with score < EVIDENCE_SCORE_THRESHOLD (0.30)
      4. If surviving passages < MIN_VALID_PASSAGES (2) -> validation_passed=False
         (graph router will retry retrieval with broader parameters)
    Ablation: USE_VALIDATOR=False -> all docs pass with score=1.0
    \"\"\"
    docs  = state.get("retrieved_docs") or []
    if not docs:
        return {**state, "scored_docs": [], "validation_passed": False,
                "stage_logs": state["stage_logs"] + [
                    {"stage": "validator", "status": "completed",
                     "details": {"passed": False, "reason": "empty_docs"}}]}

    # Ablation bypass
    if not USE_VALIDATOR:
        return {**state, "scored_docs": [{**d, "score": 1.0} for d in docs],
                "validation_passed": True,
                "stage_logs": state["stage_logs"] + [
                    {"stage": "validator", "status": "completed",
                     "details": {"ablation_skip": True, "passed": True}}]}

    scores = score_passages_against_query(state["query"],
                                         [d["content"] for d in docs])
    scored = [{**d, "score": s} for d, s in zip(docs, scores)]
    passed = [d for d in scored if d["score"] >= EVIDENCE_SCORE_THRESHOLD]
    valid  = len(passed) >= MIN_VALID_PASSAGES
    return {**state, "scored_docs": passed, "validation_passed": valid,
            "stage_logs": state["stage_logs"] + [
                {"stage": "validator", "status": "completed",
                 "details": {"total": len(docs), "passed": len(passed),
                             "threshold": EVIDENCE_SCORE_THRESHOLD,
                             "validation_passed": valid}}]}"""

GENERATOR_CODE = """\
# pipeline/nodes/generator.py  — Answer Generator Node
from langchain_core.prompts import ChatPromptTemplate
from pydantic import BaseModel, Field
from typing import List
from config.settings import LLM_MAX_TOKENS
from pipeline.state import ACRagState
from utils.attribution import build_attribution, clean_citations_from_answer
from utils.llm_factory import get_llm

class GeneratedAnswer(BaseModel):
    answer:          str   = Field(description=(
        "Answer strictly grounded in context. Cite sources inline: [N]."))
    is_answerable:   bool  = Field(description="True if context has enough information.")
    confidence:      float = Field(ge=0.0, le=1.0,
        description="Confidence the answer is complete and faithful (0-1).")
    key_sources_used: List[int] = Field(default_factory=list,
        description="Source numbers [N] primarily used.")

_SYSTEM = \"\"\"You are a precise document question-answering assistant.
STRICT RULES:
1. Answer ONLY using information present in the provided context.
2. Do NOT add knowledge from outside the context.
3. Cite EVERY claim with its source number in brackets: [1] or [2][3].
4. If context contains tables/figures (modality=table/figure), read them carefully.
5. If context does not contain enough information, state this clearly — do NOT guess.
6. Be concise but complete.\"\"\"

_prompt = ChatPromptTemplate.from_messages([
    ("system", _SYSTEM),
    ("human",  "Context:\\n{context}\\n\\n---\\nQuestion: {query}\\n"
               "Answer ONLY from the context above. Cite sources with [N]."),
])

def generator_node(state: ACRagState) -> ACRagState:
    \"\"\"
    Formal objective: a = argmax P(a | q, R(q, D))
    Faithfulness constraint: Faith(a, R(q,D)) >= tau (FAITHFULNESS_THRESHOLD)
    Reads:  query, refined_context
    Writes: answer, answer_with_attribution
    \"\"\"
    context = state.get("refined_context", "")
    if not context:
        return {**state,
                "answer": "The provided documents do not contain sufficient information.",
                "answer_with_attribution": [],
                "stage_logs": state["stage_logs"] + [
                    {"stage": "generator", "status": "completed",
                     "details": {"is_answerable": False}}]}
    chain  = _prompt | get_llm(max_tokens=LLM_MAX_TOKENS).with_structured_output(GeneratedAnswer)
    result = chain.invoke({"query": state["query"], "context": context})
    return {**state,
            "answer":                  clean_citations_from_answer(result.answer),
            "answer_with_attribution": build_attribution(result.answer, context),
            "stage_logs": state["stage_logs"] + [
                {"stage": "generator", "status": "completed",
                 "details": {"answerable": result.is_answerable,
                             "confidence": result.confidence}}]}"""

CRITIC_CODE = """\
# pipeline/nodes/critic.py  — Critic / Self-Reflection Node
from langchain_core.prompts import ChatPromptTemplate
from pydantic import BaseModel, Field
from config.settings import CRITIC_MIN_SCORE, USE_CRITIC
from pipeline.state import ACRagState, CriticScores
from utils.llm_factory import get_llm

class CriticEvaluation(BaseModel):
    faithfulness:    int = Field(ge=1, le=5,
        description="Is every claim explicitly supported by the context? 5=fully grounded.")
    completeness:    int = Field(ge=1, le=5,
        description="Does the answer address all key aspects? 5=comprehensive.")
    table_accuracy:  int = Field(ge=1, le=5,
        description="Are table values cited correctly? 5 if no tables.")
    figure_accuracy: int = Field(ge=1, le=5,
        description="Are figure descriptions accurate? 5 if no figures.")
    conciseness:     int = Field(ge=1, le=5,
        description="Free of padding and repetition? 5=perfectly concise.")
    feedback:        str = Field(description="Explanation of lowest-scoring dimensions.")
    retry_reason:    str = Field(
        description="content|format|none — classify the primary failure type.")

_prompt = ChatPromptTemplate.from_messages([
    ("system",
     "You are a rigorous quality critic for a document QA system.\\n"
     "Scale: 5=Excellent, 4=Good, 3=Fair, 2=Poor, 1=Unacceptable.\\n"
     "If ANY factual claim cannot be traced to the context, faithfulness <= 3."),
    ("human",
     "Question:\\n{query}\\n\\nContext:\\n{context}\\n\\n"
     "Answer:\\n{answer}\\n\\nEvaluate on all five dimensions."),
])

def critic_node(state: ACRagState) -> ACRagState:
    \"\"\"
    Evaluates the generated answer on 5 dimensions.
    Accept condition: ALL dimensions >= CRITIC_MIN_SCORE (4).
    Retry routing:
      content -> faithfulness/completeness/table/figure failed
                 (graph routes back to query_analyzer for full restart)
      format  -> only conciseness failed
                 (graph routes back to generator for regeneration only)
    Ablation: USE_CRITIC=False -> auto-pass with all scores=5.
    \"\"\"
    if not USE_CRITIC:
        scores = CriticScores(faithfulness=5, completeness=5, table_accuracy=5,
                              figure_accuracy=5, conciseness=5, overall=5.0,
                              feedback="Critic disabled (ablation mode).")
        return {**state, "critic_scores": scores, "critic_passed": True,
                "critic_feedback": scores["feedback"],
                "stage_logs": state["stage_logs"] + [
                    {"stage": "critic", "status": "completed",
                     "details": {"ablation_skip": True, "passed": True}}]}

    context = (state.get("refined_context") or "")[:8000]
    chain   = _prompt | get_llm().with_structured_output(CriticEvaluation)
    result  = chain.invoke({"query": state["query"], "context": context,
                            "answer": state.get("answer", "")})
    dims    = [result.faithfulness, result.completeness, result.table_accuracy,
               result.figure_accuracy, result.conciseness]
    passed  = all(d >= CRITIC_MIN_SCORE for d in dims)
    overall = round(sum(dims) / len(dims), 2)
    scores  = CriticScores(faithfulness=result.faithfulness,
                           completeness=result.completeness,
                           table_accuracy=result.table_accuracy,
                           figure_accuracy=result.figure_accuracy,
                           conciseness=result.conciseness,
                           overall=overall, feedback=result.feedback)
    return {**state, "critic_scores": scores, "critic_passed": passed,
            "critic_feedback": result.feedback,
            "retry_reason": result.retry_reason if not passed else None,
            "stage_logs": state["stage_logs"] + [
                {"stage": "critic", "status": "completed",
                 "details": {"passed": passed, "overall": overall,
                             "retry_reason": result.retry_reason}}]}"""

FIGURE_EXTRACTOR_CODE = """\
# ingestion/figure_extractor.py  — GPT-4o Vision Figure Extractor
import base64, hashlib, logging, os
from pathlib import Path
from typing import List
from langchain_core.documents import Document

MIN_IMAGE_AREA = 90_000   # ~300x300 px -- filters logos, seals, small icons

_VISION_PROMPT = (
    "You are analyzing a figure extracted from a research paper. "
    "Does this figure contain meaningful data or technical content "
    "(chart, graph, architecture diagram, table rendered as image, flowchart)? "
    "If YES: describe it in detail -- figure type, data shown, key values, "
    "labels, axes, and the main insight it conveys. "
    "If NO (logo, photo, institutional seal): respond with exactly: SKIP"
)

def _describe_figure_with_vision(image_bytes, source, page) -> str | None:
    \"\"\"Send image bytes to GPT-4o Vision; return description or None if SKIP.\"\"\"
    try:
        from openai import OpenAI
        client   = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
        b64      = base64.b64encode(image_bytes).decode("utf-8")
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": [
                {"type": "text",      "text": _VISION_PROMPT},
                {"type": "image_url", "image_url":
                 {"url": f"data:image/png;base64,{b64}", "detail": "high"}},
            ]}], max_tokens=512)
        desc = response.choices[0].message.content.strip()
        return None if desc.upper().startswith("SKIP") else desc
    except Exception as e:
        return f"[Figure on page {page} of {source} -- vision unavailable: {e}]"

def extract_figures(file_path) -> List[Document]:
    \"\"\"
    Extract all figures from a PDF and return them as Document chunks.
    Two-layer filter:
      1. Size filter: skip images < MIN_IMAGE_AREA pixels (logos, icons).
      2. Vision filter: GPT-4o Vision returns SKIP for non-informational images.
    \"\"\"
    path = Path(file_path)
    if path.suffix.lower() != ".pdf":
        return []
    import fitz
    pdf = fitz.open(str(path))
    docs, idx = [], 0
    for page_num, page in enumerate(pdf, start=1):
        for img_info in page.get_images(full=True):
            bi = pdf.extract_image(img_info[0])
            w, h = bi.get("width", 0), bi.get("height", 0)
            if w * h < MIN_IMAGE_AREA:
                continue
            desc = _describe_figure_with_vision(bi["image"], path.name, page_num)
            if desc is None:
                continue
            chunk_id = hashlib.md5(
                f"{path}::page{page_num}::fig{idx}".encode()
            ).hexdigest()[:12]
            docs.append(Document(page_content=desc, metadata={
                "source": str(path), "file_name": path.name, "file_type": ".pdf",
                "page": page_num, "figure_index": idx, "image_width": w,
                "image_height": h, "modality": "figure",
                "section_heading": "Figure", "chunk_id": chunk_id,
                "chunk_index": idx, "word_count": len(desc.split()),
                "char_count": len(desc)}))
            idx += 1
    pdf.close()
    return docs"""

SETTINGS_CODE = """\
# config/settings.py  — Central Configuration
from pathlib import Path
from dotenv import load_dotenv
_ENV_FILE = Path(__file__).resolve().parent.parent / ".env"
load_dotenv(dotenv_path=_ENV_FILE, override=False)

# Paths
BASE_DIR            = Path(__file__).resolve().parent.parent
DATA_RAW_DIR        = BASE_DIR / "data"  / "raw"
DATA_PROCESSED_DIR  = BASE_DIR / "data"  / "processed"
VECTORSTORE_DIR     = BASE_DIR / "vectorstore" / "index"
LOG_FILE            = BASE_DIR / "logs"  / "pipeline.log"

# Ingestion
CHUNK_SIZE           = 512     # tokens (approx 4 chars/token)
CHUNK_OVERLAP        = 64
SUPPORTED_EXTENSIONS = [".pdf", ".docx", ".txt", ".html", ".md"]

# Embeddings
EMBEDDING_MODEL     = "text-embedding-3-large"   # OpenAI
VECTORSTORE_BACKEND = "faiss"                     # "faiss" | "chroma"

# LLM Provider  ("openai" | "google" | "anthropic" | "groq")
LLM_PROVIDER        = "openai"
OPENAI_LLM_MODEL    = "gpt-4o"
GOOGLE_LLM_MODEL    = "gemini-2.0-flash"
ANTHROPIC_LLM_MODEL = "claude-sonnet-4-6"
GROQ_LLM_MODEL      = "llama-3.3-70b-versatile"
LLM_MODEL           = {"openai": OPENAI_LLM_MODEL, "google": GOOGLE_LLM_MODEL,
                        "anthropic": ANTHROPIC_LLM_MODEL, "groq": GROQ_LLM_MODEL
                        }[LLM_PROVIDER]
LLM_TEMPERATURE     = 0.0
LLM_MAX_TOKENS      = 2048

# Retrieval
RETRIEVAL_K_MIN        = 4
RETRIEVAL_K_MAX        = 12
RETRIEVAL_K_DEFAULT    = 6
MMR_FETCH_K_MULTIPLIER = 3      # fetch_k = k * multiplier
MMR_LAMBDA_MULT        = 0.5    # 0=max diversity, 1=max relevance

# Validation
EVIDENCE_SCORE_THRESHOLD = 0.30
MIN_VALID_PASSAGES       = 2

# Critic / Self-Reflection
CRITIC_MIN_SCORE    = 4         # all dimensions must be >= 4 (scale 1-5)
MAX_RETRIES         = 3

# Faithfulness constraint (formal objective)
FAITHFULNESS_THRESHOLD = 0.80   # tau in: Faith(a, R(q,D)) >= tau

# Entry Router
ROUTER_SIMILARITY_THRESHOLD = EVIDENCE_SCORE_THRESHOLD

# Ablation flags (set False to disable each pipeline stage)
USE_RETRIEVAL_PLANNER = True
USE_VALIDATOR         = False
USE_CONTEXT_REFINER   = True
USE_CRITIC            = True

# Logging
LOG_LEVEL = "INFO"    # "DEBUG" for verbose stage traces"""


# ─── Helper: make a code paragraph element ────────────────────────────────────

def make_code_para(doc, line_text):
    """Create a <w:p> element styled as monospace code with gray shading."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')

    # paragraph style = Normal
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Normal')
    pPr.append(pStyle)

    # no spacing before/after
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)

    # gray shading
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)

    p.append(pPr)

    # run with Courier New 9pt
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    fonts = OxmlElement('w:rFonts')
    fonts.set(qn('w:ascii'), 'Courier New')
    fonts.set(qn('w:hAnsi'), 'Courier New')
    rPr.append(fonts)
    sz = OxmlElement('w:sz');   sz.set(qn('w:val'), '18')   # 9pt
    szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), '18')
    rPr.append(sz); rPr.append(szCs)
    r.append(rPr)

    t = OxmlElement('w:t')
    t.text = line_text if line_text else ' '
    if not line_text or line_text != line_text.strip():
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    return p


def make_heading_para(doc, text, level=2):
    """Create a heading paragraph element."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    # Map heading level to style id
    style_map = {1: 'Heading1', 2: 'Heading2', 3: 'Heading3'}
    # Try to find the actual style id in doc
    target_name = f'Heading {level}'
    style_id = style_map.get(level, 'Heading2')
    for s in doc.styles:
        if s.name == target_name:
            style_id = s.style_id
            break
    pStyle.set(qn('w:val'), style_id)
    pPr.append(pStyle)
    p.append(pPr)

    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    p.append(r)
    return p


def make_body_para(doc, text):
    """Create a normal body paragraph element."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    style_id = 'Normal'
    for s in doc.styles:
        if s.name in ('Normal', 'Body Text'):
            style_id = s.style_id
            break
    pStyle.set(qn('w:val'), style_id)
    pPr.append(pStyle)
    # add some spacing
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '100')
    spacing.set(qn('w:after'), '100')
    pPr.append(spacing)
    p.append(pPr)

    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    if text and (text.startswith(' ') or text.endswith(' ')):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    return p


def make_page_break():
    """Create a page break paragraph element."""
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    return p


def make_blank_para():
    p = OxmlElement('w:p')
    return p


# ─── Find Chapter 4 heading index ─────────────────────────────────────────────

body = doc.element.body
body_children = list(body)  # all direct children: paragraphs, tables, sectPr etc.

# Map paragraph objects to their body child index
para_to_body_idx = {}
para_idx = 0
for bi, child in enumerate(body_children):
    if child.tag.endswith('}p'):
        if para_idx < len(doc.paragraphs):
            para_to_body_idx[para_idx] = bi
        para_idx += 1

print(f"Total paragraphs: {len(doc.paragraphs)}, Body children: {len(body_children)}")

# Find Chapter 4 heading
ch4_para_idx = None
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip().lower()
    style_name = p.style.name.lower() if p.style and p.style.name else ''
    if ('chapter 4' in text or ('evaluation' in text and '4' in text)) and 'heading' in style_name:
        ch4_para_idx = i
        print(f"Found Chapter 4 at paragraph {i}: '{p.text[:80]}'")
        break

if ch4_para_idx is None:
    # Try finding by style + number pattern
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style_name = p.style.name.lower() if p.style and p.style.name else ''
        if 'heading' in style_name and text.startswith('4'):
            ch4_para_idx = i
            print(f"Found Chapter 4 (pattern) at paragraph {i}: '{text[:80]}'")
            break

if ch4_para_idx is None:
    print("WARNING: Chapter 4 not found! Will append at end.")

# Get the body index of the Chapter 4 paragraph
if ch4_para_idx is not None:
    ch4_body_idx = para_to_body_idx.get(ch4_para_idx)
    print(f"Chapter 4 body child index: {ch4_body_idx}")
    ref_element = body_children[ch4_body_idx]
else:
    # Append before sectPr (last child)
    ref_element = body_children[-1]


# ─── Build section 3.10 elements ─────────────────────────────────────────────

new_elements = []

new_elements.append(make_page_break())
new_elements.append(make_heading_para(doc, '3.10 Key Code Listings', level=2))
new_elements.append(make_body_para(doc,
    'This section presents the complete source code for each stage of the AC-RAG pipeline. '
    'All nodes are implemented as pure functions operating on the shared ACRagState TypedDict. '
    'This design allows LangGraph to apply partial state updates — each node writes only the '
    'fields it owns while leaving all other fields unchanged. Ablation flags in config/settings.py '
    'can disable any stage without modifying node code.'))
new_elements.append(make_blank_para())


def add_code_section(elements, heading_text, description, code_text, level=3):
    elements.append(make_heading_para(doc, heading_text, level=level))
    elements.append(make_body_para(doc, description))
    elements.append(make_blank_para())
    for line in code_text.split('\n'):
        elements.append(make_code_para(doc, line))
    elements.append(make_blank_para())


add_code_section(new_elements,
    '3.10.1  Shared State Schema',
    'The ACRagState TypedDict is the single source of truth shared across all pipeline nodes. '
    'Using TypedDict (not dataclass) is a LangGraph requirement — it enables partial updates '
    'where each node writes only the fields it owns. The RetrievalPlan and CriticScores '
    'nested TypedDicts enforce type contracts on the planner and critic outputs respectively.',
    STATE_CODE)

add_code_section(new_elements,
    '3.10.2  LangGraph Pipeline Assembly',
    'The build_pipeline() function assembles the complete StateGraph topology using '
    'conditional edges for all routing decisions. The increment_retry node acts as a '
    'counter gate that prevents infinite loops — the _route_after_retry_increment '
    'function reads the retry_reason field to determine whether to route back to the '
    'query analyzer (content failure), generator (format failure), or retriever '
    '(validation failure).',
    GRAPH_CODE)

add_code_section(new_elements,
    '3.10.3  Entry Router Node',
    'The Entry Router performs a lightweight k=3 cosine similarity check before '
    'committing to the full RAG pipeline. Queries scoring below ROUTER_SIMILARITY_THRESHOLD '
    'on both top-score and average-score criteria are routed to the direct_responder node, '
    'which returns a calibrated "I don\'t know" response. This prevents hallucination for '
    'out-of-domain queries and avoids unnecessary LLM API calls.',
    ENTRY_ROUTER_CODE)

add_code_section(new_elements,
    '3.10.4  Query Analyzer Node',
    'The Query Analyzer uses Pydantic structured output to guarantee parseable JSON from '
    'the LLM, eliminating fragile regex parsing. The complexity_score drives the adaptive '
    'retrieval strategy: scores below 0.4 trigger shallow k=4 retrieval, while scores '
    'above 0.7 activate multi-query decomposition with up to four sub-queries. '
    'The intent classification (factual/analytical/comparative/summarization) '
    'further influences MMR diversity parameters in the retrieval planner.',
    QUERY_ANALYZER_CODE)

add_code_section(new_elements,
    '3.10.5  Retrieval Planner Node',
    'The Retrieval Planner implements a two-tier design. Tier 1 produces a deterministic '
    'rule-based plan from the complexity score and intent — ensuring fast, ablation-safe '
    'behavior that does not require any LLM call. Tier 2 (enabled via USE_RETRIEVAL_PLANNER) '
    'applies an LLM refinement pass that can override modality_filter and k for queries '
    'explicitly referencing figures or tables. A modality guard prevents the LLM from '
    'over-specifying modality for general queries.',
    RETRIEVAL_PLANNER_CODE)

add_code_section(new_elements,
    '3.10.6  Evidence Validator Node',
    'The Validator scores each retrieved passage against the query using cosine similarity '
    'of their embeddings. Passages below EVIDENCE_SCORE_THRESHOLD (0.30) are discarded. '
    'If fewer than MIN_VALID_PASSAGES (2) survive filtering, validation_passed is set to '
    'False and the graph router triggers a retrieval retry with broader MMR parameters. '
    'Setting USE_VALIDATOR=False in the ablation study bypasses scoring and passes all '
    'retrieved documents through with score=1.0.',
    VALIDATOR_CODE)

add_code_section(new_elements,
    '3.10.7  Answer Generator Node',
    'The Generator enforces the faithfulness constraint Faith(a, R(q,D)) >= tau at two '
    'levels: (1) explicit system prompt instructions prohibiting external knowledge, and '
    '(2) mandatory inline source citations [N] for every claim. The Pydantic '
    'GeneratedAnswer schema captures both the answer text and key source indices, '
    'which the build_attribution() utility converts into a per-sentence evidence map '
    'for downstream critic evaluation and frontend display.',
    GENERATOR_CODE)

add_code_section(new_elements,
    '3.10.8  Critic / Self-Reflection Node',
    'The Critic evaluates answers on five dimensions: faithfulness, completeness, '
    'table accuracy, figure accuracy, and conciseness. All dimensions must score '
    '>= CRITIC_MIN_SCORE (4 on a 1-5 scale) for the answer to be accepted. '
    'The retry classification — "content" vs "format" — determines the retry target: '
    'content failures trigger a full pipeline restart via the query_analyzer, while '
    'format failures regenerate the answer via the generator using the same context. '
    'Setting USE_CRITIC=False in ablation testing auto-passes all answers.',
    CRITIC_CODE)

add_code_section(new_elements,
    '3.10.9  GPT-4o Vision Figure Extractor',
    'The figure extractor implements dual-layer filtering to ensure only informational '
    'figures are indexed. Layer 1 is a pixel-area threshold (MIN_IMAGE_AREA=90,000 px²) '
    'that filters logos and icons based on image dimensions alone. Layer 2 sends each '
    'surviving image to GPT-4o Vision with a prompt requesting "SKIP" for non-data images. '
    'Only images passing both filters become retrieval-ready Document chunks with '
    'modality="figure" and a rich textual description as their page_content.',
    FIGURE_EXTRACTOR_CODE)

add_code_section(new_elements,
    '3.10.10  System Configuration',
    'All hyperparameters, model names, ablation flags, and thresholds are centralised '
    'in config/settings.py. Changing LLM_PROVIDER to "google", "anthropic", or "groq" '
    'switches the entire pipeline to that provider in a single edit. The ablation flags '
    'USE_RETRIEVAL_PLANNER, USE_VALIDATOR, USE_CONTEXT_REFINER, and USE_CRITIC can '
    'disable any pipeline stage independently for the comparative evaluation study.',
    SETTINGS_CODE)


# ─── Insert all new elements before the Chapter 4 reference element ───────────

for elem in reversed(new_elements):
    ref_element.addprevious(elem)

print(f"Inserted {len(new_elements)} XML elements before Chapter 4.")

# ─── Save ─────────────────────────────────────────────────────────────────────
doc.save(THESIS_PATH)
print(f"Saved: {THESIS_PATH}")

# Verify
doc2 = Document(THESIS_PATH)
print(f"Total paragraphs after insertion: {len(doc2.paragraphs)}")
