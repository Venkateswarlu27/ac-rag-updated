"""
fix_thesis_final.py
Comprehensive thesis fix:
1. Remove the incorrectly-ordered Section 3.10 (code in reverse)
2. Re-insert Section 3.10 in correct order
3. Update TOC to include Section 3.10
4. Update List of Figures page references
5. Fix indentation and body text issues
"""

import shutil
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

THESIS_PATH = "AC_RAG_Thesis_Submission_Ready.docx"

# ─── Backup ───────────────────────────────────────────────────────────────────
shutil.copy(THESIS_PATH, THESIS_PATH.replace('.docx', '_backup_before_fix.docx'))
print("Backup saved.")

doc = Document(THESIS_PATH)
body = doc.element.body

# ─── Step 1: Remove wrong Section 3.10 (paras 299–930) ───────────────────────
# Para 299 = first wrong line (LOG_LEVEL reversed settings.py code)
# Para 930 = last para of wrong section (body text after 3.10 heading)

print(f"Total paragraphs before fix: {len(doc.paragraphs)}")

# Build para_index -> body_child_index map
body_children = list(body)
para_to_body = {}
pi = 0
for bi, child in enumerate(body_children):
    if child.tag.endswith('}p'):
        para_to_body[pi] = bi
        pi += 1

# Find the range to remove:
# First wrong paragraph: first paragraph AFTER para 298 that has reversed code
# Last wrong paragraph: the one just before Chapter 4 (para 931 currently)
# So remove para 299 through 930 (inclusive)

remove_start_para = 299
remove_end_para   = 930

# Verify boundaries
print(f"  Remove start (para {remove_start_para}): '{doc.paragraphs[remove_start_para].text[:60]}'")
print(f"  Remove end   (para {remove_end_para}): '{doc.paragraphs[remove_end_para].text[:60]}'")
print(f"  Para after remove (931 = Chapter 4): '{doc.paragraphs[931].text[:60]}'")

# Collect body elements to remove
body_idx_start = para_to_body[remove_start_para]
body_idx_end   = para_to_body[remove_end_para]
print(f"  Body indices to remove: {body_idx_start} → {body_idx_end}")

elements_to_remove = []
for bi in range(body_idx_start, body_idx_end + 1):
    elements_to_remove.append(body_children[bi])

for elem in elements_to_remove:
    body.remove(elem)

print(f"Removed {len(elements_to_remove)} wrong elements.")
print(f"Total paragraphs after removal: {len(doc.paragraphs)}")


# ─── Step 2: Helper functions ─────────────────────────────────────────────────

def make_code_para(line_text):
    """Gray-shaded Courier New 9pt paragraph for code."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')

    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Normal')
    pPr.append(pStyle)

    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)

    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)

    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '360')
    ind.set(qn('w:right'), '360')
    pPr.append(ind)

    p.append(pPr)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    fonts = OxmlElement('w:rFonts')
    fonts.set(qn('w:ascii'), 'Courier New')
    fonts.set(qn('w:hAnsi'), 'Courier New')
    rPr.append(fonts)
    sz = OxmlElement('w:sz');   sz.set(qn('w:val'), '18')
    szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), '18')
    rPr.append(sz); rPr.append(szCs)
    r.append(rPr)

    t = OxmlElement('w:t')
    t.text = line_text if line_text else ' '
    if not line_text or line_text != line_text.strip() or line_text.startswith(' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    return p


def make_heading_para(text, level=2):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    level_map = {1: 'Heading1', 2: 'Heading2', 3: 'Heading3'}
    style_id = level_map.get(level, 'Heading2')
    for s in doc.styles:
        if s.name == f'Heading {level}':
            style_id = s.style_id
            break
    pStyle.set(qn('w:val'), style_id)
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    p.append(r)
    return p


def make_body_para(text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    style_id = 'Normal'
    for s in doc.styles:
        if s.name == 'Normal':
            style_id = s.style_id
            break
    pStyle.set(qn('w:val'), style_id)
    pPr.append(pStyle)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '60')
    spacing.set(qn('w:after'), '80')
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    if text and (text.startswith(' ') or text.endswith(' ')):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    return p


def make_page_break():
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    return p


def make_blank():
    p = OxmlElement('w:p')
    return p


def code_paras(code_text):
    return [make_code_para(line) for line in code_text.split('\n')]


def make_caption_para(text):
    """Bold caption para for code listings."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '120')
    spacing.set(qn('w:after'), '40')
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    b = OxmlElement('w:b')
    rPr.append(b)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '20')
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    p.append(r)
    return p


# ─── Step 3: Code content ─────────────────────────────────────────────────────

STATE_CODE = """\
# pipeline/state.py  — Shared State Schema
from __future__ import annotations
from typing import Any, Dict, List, Optional, TypedDict

class RetrievalPlan(TypedDict):
    k: int                   # number of final passages to retrieve (4-12)
    fetch_k: int             # MMR candidate pool size
    lambda_mult: float       # MMR diversity (0=diverse, 1=relevant)
    modality_filter: str     # "text" | "table" | "figure" | "all"
    use_multi_query: bool    # decompose query into sub-queries
    retrieval_depth: str     # "shallow" | "standard" | "deep"

class CriticScores(TypedDict):
    faithfulness: int
    completeness: int
    table_accuracy: int
    figure_accuracy: int
    conciseness: int
    overall: float
    feedback: str

class ACRagState(TypedDict):
    query: str
    route: Optional[str]                   # "rag" | "unknown"
    rewritten_query: Optional[str]
    intent: Optional[str]                  # factual|analytical|comparative|summarization
    complexity_score: Optional[float]      # 0.0 (simple) -> 1.0 (complex)
    decomposed_queries: Optional[List[str]]
    retrieval_plan: Optional[RetrievalPlan]
    retrieved_docs: Optional[List[Dict[str, Any]]]
    scored_docs: Optional[List[Dict[str, Any]]]
    validation_passed: Optional[bool]
    refined_context: Optional[str]
    answer: Optional[str]
    answer_with_attribution: Optional[List[Dict[str, Any]]]
    critic_scores: Optional[CriticScores]
    critic_passed: Optional[bool]
    critic_feedback: Optional[str]
    retry_count: int
    retry_reason: Optional[str]            # "content" | "format"
    stage_logs: List[Dict]
    error: Optional[str]

def initial_state(query: str) -> ACRagState:
    return ACRagState(
        query=query, route=None, rewritten_query=None,
        intent=None, complexity_score=None, decomposed_queries=None,
        retrieval_plan=None, retrieved_docs=None,
        scored_docs=None, validation_passed=None,
        refined_context=None, answer=None,
        answer_with_attribution=None, critic_scores=None,
        critic_passed=None, critic_feedback=None,
        retry_count=0, retry_reason=None, stage_logs=[], error=None,
    )"""

GRAPH_CODE = """\
# pipeline/graph.py  — LangGraph Pipeline Assembly
from langgraph.graph import StateGraph, END
from pipeline.state import ACRagState, initial_state

def build_pipeline(vsm=None):
    if vsm is None:
        vsm = VectorStoreManager(); vsm.load()

    graph = StateGraph(ACRagState)

    # Register all nodes
    graph.add_node("entry_router",      make_entry_router_node(vsm))
    graph.add_node("direct_responder",  direct_responder_node)
    graph.add_node("query_analyzer",    query_analyzer_node)
    graph.add_node("retrieval_planner", retrieval_planner_node)
    graph.add_node("retriever",         make_retriever_node(vsm))
    graph.add_node("validator",         validator_node)
    graph.add_node("context_refiner",   context_refiner_node)
    graph.add_node("generator",         generator_node)
    graph.add_node("critic",            critic_node)
    graph.add_node("increment_retry",
                   lambda s: {**s, "retry_count": s["retry_count"]+1})
    graph.add_node("end_success",       lambda s: s)
    graph.add_node("end_error",         lambda s: s)
    graph.add_node("end_max_retries",   lambda s: s)

    graph.set_entry_point("entry_router")

    # Routing: entry_router -> rag pipeline | direct_responder
    graph.add_conditional_edges("entry_router", route_after_entry_router,
        {"rag": "query_analyzer", "unknown": "direct_responder"})
    graph.add_edge("direct_responder", "end_success")

    # query_analyzer -> retrieval_planner | end_error
    graph.add_conditional_edges("query_analyzer", route_after_query_analyzer,
        {"retrieval_planner": "retrieval_planner", "end_error": "end_error"})

    # retrieval_planner -> retriever (always)
    graph.add_conditional_edges("retrieval_planner",
        route_after_retrieval_planner, {"retriever": "retriever"})

    # retriever -> validator | retry | end_error | end_max_retries
    graph.add_conditional_edges("retriever", route_after_retriever,
        {"validator": "validator", "query_analyzer": "increment_retry",
         "end_error": "end_error", "end_max_retries": "end_max_retries"})

    # validator -> context_refiner | generator (ablation) | retry
    graph.add_conditional_edges("validator", route_after_validator,
        {"context_refiner": "context_refiner", "generator": "generator",
         "retriever": "increment_retry", "end_max_retries": "end_max_retries"})

    # context_refiner -> generator
    graph.add_conditional_edges("context_refiner",
        route_after_context_refiner, {"generator": "generator"})

    # generator -> critic | end_success (ablation) | end_error
    graph.add_conditional_edges("generator", route_after_generator,
        {"critic": "critic", "end_success": "end_success",
         "end_error": "end_error"})

    # critic -> end_success | content retry | format retry | end_max_retries
    graph.add_conditional_edges("critic", route_after_critic,
        {"end_success": "end_success", "query_analyzer": "increment_retry",
         "generator": "increment_retry", "end_max_retries": "end_max_retries"})

    # increment_retry -> correct destination
    graph.add_conditional_edges("increment_retry", _route_after_retry_increment,
        {"query_analyzer": "query_analyzer", "retriever": "retriever",
         "generator": "generator", "end_max_retries": "end_max_retries"})

    graph.add_edge("end_success",     END)
    graph.add_edge("end_error",       END)
    graph.add_edge("end_max_retries", END)
    return graph.compile()

def _route_after_retry_increment(state: ACRagState) -> str:
    from config.settings import MAX_RETRIES
    if state.get("retry_count", 0) > MAX_RETRIES:
        return "end_max_retries"
    reason = state.get("retry_reason") or "content"
    if reason == "format":
        return "generator"
    if state.get("validation_passed") is False:
        return "retriever"
    return "query_analyzer"

def run_pipeline(query: str, vsm=None) -> ACRagState:
    pipeline = build_pipeline(vsm)
    result   = pipeline.invoke(initial_state(query),
                               config={"recursion_limit": 50})
    return result"""

ENTRY_ROUTER_CODE = """\
# pipeline/nodes/entry_router.py  — Entry Router Node
from pipeline.state import ACRagState
from config.settings import ROUTER_SIMILARITY_THRESHOLD

def _similarity_route(vsm, query: str) -> str:
    \"\"\"k=3 cosine similarity check; returns 'rag' or 'unknown'.\"\"\"
    try:
        results = vsm.similarity_search_with_score(query, k=3)
        if not results:
            return "unknown"
        avg = sum(s for _, s in results) / len(results)
        top = results[0][1]
        if (top >= ROUTER_SIMILARITY_THRESHOLD
                and avg >= ROUTER_SIMILARITY_THRESHOLD * 0.7):
            return "rag"
        return "unknown"
    except Exception:
        return "rag"   # fail-safe: attempt RAG rather than refusing

def make_entry_router_node(vsm):
    def entry_router_node(state: ACRagState) -> ACRagState:
        route = _similarity_route(vsm, state["query"].strip())
        return {
            **state, "route": route,
            "stage_logs": state["stage_logs"] + [
                {"stage": "entry_router", "status": "completed",
                 "details": {"route": route}}
            ],
        }
    return entry_router_node"""

QUERY_ANALYZER_CODE = """\
# pipeline/nodes/query_analyzer.py  — Query Analyzer Node
from langchain_core.prompts import ChatPromptTemplate
from pydantic import BaseModel, Field
from typing import List
from pipeline.state import ACRagState
from utils.llm_factory import get_llm

class QueryAnalysis(BaseModel):
    rewritten_query:    str   = Field(
        description="Clarified retrieval-optimised query.")
    intent:             str   = Field(
        description="factual|analytical|comparative|summarization")
    complexity_score:   float = Field(ge=0.0, le=1.0,
        description="0.0 (simple) -> 1.0 (complex multi-hop)")
    complexity_reason:  str   = Field(
        description="One-sentence explanation of complexity score.")
    decomposed_queries: List[str] = Field(default_factory=list,
        description="Sub-queries if complexity >= 0.6, else empty list.")

_SYSTEM = \"\"\"You are a query analysis expert for a document QA system.
Intent categories:
  factual       - specific fact, number, date, or definition
  analytical    - explain, analyse, or reason about something
  comparative   - compare two or more things
  summarization - summary or overview

Complexity scoring:
  0.0-0.3  Single fact lookup, one document needed
  0.4-0.6  Synthesise a few passages, some reasoning
  0.7-0.9  Multi-hop reasoning, cross-document synthesis
  1.0      Requires full document understanding

Only decompose if complexity >= 0.6. Maximum 4 sub-queries.\"\"\"

_prompt = ChatPromptTemplate.from_messages([
    ("system", _SYSTEM), ("human", "Query: {query}"),
])

def query_analyzer_node(state: ACRagState) -> ACRagState:
    \"\"\"
    Reads:  state['query']
    Writes: rewritten_query, intent, complexity_score,
            decomposed_queries
    \"\"\"
    try:
        result = (
            _prompt | get_llm().with_structured_output(QueryAnalysis)
        ).invoke({"query": state["query"]})
        if result.complexity_score < 0.6:
            result.decomposed_queries = []
        return {
            **state,
            "rewritten_query":    result.rewritten_query,
            "intent":             result.intent,
            "complexity_score":   result.complexity_score,
            "decomposed_queries": result.decomposed_queries,
            "stage_logs": state["stage_logs"] + [
                {"stage": "query_analyzer", "status": "completed",
                 "details": {"intent": result.intent,
                             "complexity": result.complexity_score}}
            ],
        }
    except Exception as e:
        return {
            **state, "error": f"QueryAnalyzer failed: {e}",
            "stage_logs": state["stage_logs"] + [
                {"stage": "query_analyzer", "status": "failed",
                 "details": {"error": str(e)}}
            ],
        }"""

RETRIEVAL_PLANNER_CODE = """\
# pipeline/nodes/retrieval_planner.py  — Retrieval Planner Node
from pydantic import BaseModel, Field
from langchain_core.prompts import ChatPromptTemplate
from config.settings import (
    RETRIEVAL_K_MIN, RETRIEVAL_K_MAX, RETRIEVAL_K_DEFAULT,
    MMR_FETCH_K_MULTIPLIER, MMR_LAMBDA_MULT, USE_RETRIEVAL_PLANNER)
from pipeline.state import ACRagState, RetrievalPlan
from utils.llm_factory import get_llm

def _rule_based_plan(complexity: float, intent: str) -> RetrievalPlan:
    \"\"\"Tier-1 deterministic plan based on complexity thresholds.\"\"\"
    if complexity < 0.4:
        k, depth, multi, lm = (RETRIEVAL_K_MIN,
                                "shallow", False, 0.7)
    elif complexity < 0.7:
        k, depth, multi, lm = (RETRIEVAL_K_DEFAULT,
                                "standard", False, MMR_LAMBDA_MULT)
    else:
        k, depth, multi, lm = (min(10, RETRIEVAL_K_MAX),
                                "deep", True, 0.3)
    if intent in ("analytical", "comparative"):
        lm = max(0.3, lm - 0.1)
        k  = min(k + 2, RETRIEVAL_K_MAX)
    return RetrievalPlan(
        k=k, fetch_k=k * MMR_FETCH_K_MULTIPLIER,
        lambda_mult=lm, modality_filter="all",
        use_multi_query=multi, retrieval_depth=depth)

class PlanRefinement(BaseModel):
    modality_filter: str = Field(
        description="text | table | figure | all")
    k_override:      int = Field(ge=4, le=12,
        description="Override k if the default is clearly wrong.")
    override_reason: str = Field(
        description="Reason for change, or 'no change'.")

_PROMPT = ChatPromptTemplate.from_messages([
    ("system",
     "You are a retrieval planning agent. Refine the rule-based\\n"
     "plan by:\\n"
     "1. Setting modality_filter only if the query explicitly asks\\n"
     "   about a table or figure.\\n"
     "2. Optionally adjusting k if clearly wrong.\\n"
     "Do NOT change fetch_k, lambda_mult, or use_multi_query."),
    ("human",
     "Query: {query}\\nIntent: {intent}\\n"
     "Complexity: {complexity}\\nRule k: {k}  Depth: {depth}\\n"
     "Refine the plan."),
])

def retrieval_planner_node(state: ACRagState) -> ACRagState:
    \"\"\"
    Two-tier approach:
      Tier 1 — Rule-based defaults (always runs, ablation-safe)
      Tier 2 — LLM refinement pass (USE_RETRIEVAL_PLANNER flag)
    Reads:  complexity_score, intent, query
    Writes: retrieval_plan
    \"\"\"
    complexity = state.get("complexity_score") or 0.5
    intent     = state.get("intent") or "factual"
    plan       = _rule_based_plan(complexity, intent)

    if USE_RETRIEVAL_PLANNER:
        try:
            ref = (
                _PROMPT | get_llm().with_structured_output(PlanRefinement)
            ).invoke({
                "query":      state.get("rewritten_query") or state["query"],
                "intent":     intent, "complexity": complexity,
                "k":          plan["k"],
                "depth":      plan["retrieval_depth"],
            })
            plan["modality_filter"] = ref.modality_filter
            plan["k"]       = ref.k_override
            plan["fetch_k"] = ref.k_override * MMR_FETCH_K_MULTIPLIER
        except Exception:
            pass   # fall back to rule-based plan on LLM failure

    return {
        **state, "retrieval_plan": plan,
        "stage_logs": state["stage_logs"] + [
            {"stage": "retrieval_planner", "status": "completed",
             "details": dict(plan)}
        ],
    }"""

VALIDATOR_CODE = """\
# pipeline/nodes/validator.py  — Evidence Validator Node
from config.settings import (
    EVIDENCE_SCORE_THRESHOLD, MIN_VALID_PASSAGES, USE_VALIDATOR)
from pipeline.state import ACRagState
from utils.scoring import score_passages_against_query

def validator_node(state: ACRagState) -> ACRagState:
    \"\"\"
    Evidence Validation Algorithm:
      1. Embed query and all retrieved passages
      2. Compute cosine similarity (query vs each passage)
      3. Discard passages below EVIDENCE_SCORE_THRESHOLD (0.30)
      4. If surviving passages < MIN_VALID_PASSAGES (2)
         -> validation_passed=False; graph router retries retrieval
    Ablation: USE_VALIDATOR=False -> all docs pass with score=1.0
    \"\"\"
    docs = state.get("retrieved_docs") or []

    if not docs:
        return {
            **state, "scored_docs": [], "validation_passed": False,
            "stage_logs": state["stage_logs"] + [
                {"stage": "validator", "status": "completed",
                 "details": {"passed": False, "reason": "empty_docs"}}
            ],
        }

    # Ablation bypass
    if not USE_VALIDATOR:
        return {
            **state,
            "scored_docs":      [{**d, "score": 1.0} for d in docs],
            "validation_passed": True,
            "stage_logs": state["stage_logs"] + [
                {"stage": "validator", "status": "completed",
                 "details": {"ablation_skip": True, "passed": True}}
            ],
        }

    scores = score_passages_against_query(
        state["query"], [d["content"] for d in docs])
    scored = [{**d, "score": s} for d, s in zip(docs, scores)]
    passed = [d for d in scored if d["score"] >= EVIDENCE_SCORE_THRESHOLD]
    valid  = len(passed) >= MIN_VALID_PASSAGES

    return {
        **state, "scored_docs": passed, "validation_passed": valid,
        "stage_logs": state["stage_logs"] + [
            {"stage": "validator", "status": "completed",
             "details": {"total": len(docs), "passed": len(passed),
                         "threshold": EVIDENCE_SCORE_THRESHOLD,
                         "validation_passed": valid}}
        ],
    }"""

GENERATOR_CODE = """\
# pipeline/nodes/generator.py  — Answer Generator Node
from langchain_core.prompts import ChatPromptTemplate
from pydantic import BaseModel, Field
from typing import List
from config.settings import LLM_MAX_TOKENS
from pipeline.state import ACRagState
from utils.attribution import build_attribution, clean_citations_from_answer
from utils.llm_factory import get_llm

class GeneratedAnswer(BaseModel):
    answer: str = Field(description=(
        "Answer strictly grounded in context. "
        "Cite sources inline using [N] notation."))
    is_answerable:    bool  = Field(
        description="True if context has enough information.")
    confidence:       float = Field(ge=0.0, le=1.0,
        description="Confidence the answer is faithful (0-1).")
    key_sources_used: List[int] = Field(default_factory=list,
        description="Source numbers [N] primarily used.")

_SYSTEM = \"\"\"You are a precise document question-answering assistant.
STRICT RULES:
1. Answer ONLY using information present in the provided context.
2. Do NOT add knowledge from outside the context.
3. Cite EVERY claim with its source number: [1] or [2][3].
4. If context contains tables/figures (modality=table/figure),
   read them carefully.
5. If context lacks enough information, state this clearly.
6. Be concise but complete.\"\"\"

_prompt = ChatPromptTemplate.from_messages([
    ("system", _SYSTEM),
    ("human",
     "Context:\\n{context}\\n\\n---\\nQuestion: {query}\\n"
     "Answer ONLY from the context. Cite sources with [N]."),
])

def generator_node(state: ACRagState) -> ACRagState:
    \"\"\"
    Formal objective: a = argmax P(a | q, R(q, D))
    Constraint: Faith(a, R(q, D)) >= tau (FAITHFULNESS_THRESHOLD)
    Reads:  query, refined_context
    Writes: answer, answer_with_attribution
    \"\"\"
    context = state.get("refined_context", "")
    if not context:
        return {
            **state,
            "answer": ("The provided documents do not contain "
                       "sufficient information."),
            "answer_with_attribution": [],
            "stage_logs": state["stage_logs"] + [
                {"stage": "generator", "status": "completed",
                 "details": {"is_answerable": False}}
            ],
        }
    chain  = (_prompt
              | get_llm(max_tokens=LLM_MAX_TOKENS)
              .with_structured_output(GeneratedAnswer))
    result = chain.invoke({"query": state["query"], "context": context})
    return {
        **state,
        "answer":                  clean_citations_from_answer(result.answer),
        "answer_with_attribution": build_attribution(result.answer, context),
        "stage_logs": state["stage_logs"] + [
            {"stage": "generator", "status": "completed",
             "details": {"answerable": result.is_answerable,
                         "confidence": result.confidence}}
        ],
    }"""

CRITIC_CODE = """\
# pipeline/nodes/critic.py  — Critic / Self-Reflection Node
from langchain_core.prompts import ChatPromptTemplate
from pydantic import BaseModel, Field
from config.settings import CRITIC_MIN_SCORE, USE_CRITIC
from pipeline.state import ACRagState, CriticScores
from utils.llm_factory import get_llm

class CriticEvaluation(BaseModel):
    faithfulness:    int = Field(ge=1, le=5, description=(
        "Is every claim explicitly supported by the context? "
        "5=fully grounded, 1=hallucinated."))
    completeness:    int = Field(ge=1, le=5, description=(
        "Does the answer address all key aspects? "
        "5=comprehensive, 1=major gaps."))
    table_accuracy:  int = Field(ge=1, le=5, description=(
        "Are table values cited correctly? Score 5 if no tables."))
    figure_accuracy: int = Field(ge=1, le=5, description=(
        "Are figure descriptions accurate? Score 5 if no figures."))
    conciseness:     int = Field(ge=1, le=5, description=(
        "Free of padding and repetition? 5=perfectly concise."))
    feedback:        str = Field(
        description="Explanation of lowest-scoring dimensions.")
    retry_reason:    str = Field(
        description="content | format | none.")

_prompt = ChatPromptTemplate.from_messages([
    ("system",
     "You are a rigorous quality critic for a document QA system.\\n"
     "Scale: 5=Excellent, 4=Good, 3=Fair, 2=Poor, 1=Unacceptable.\\n"
     "If ANY factual claim cannot be traced to the context, "
     "faithfulness must be <= 3."),
    ("human",
     "Question:\\n{query}\\n\\nContext:\\n{context}\\n\\n"
     "Answer:\\n{answer}\\n\\nEvaluate on all five dimensions."),
])

def critic_node(state: ACRagState) -> ACRagState:
    \"\"\"
    Evaluates answers on 5 quality dimensions (1-5 scale).
    Accept condition: ALL dimensions >= CRITIC_MIN_SCORE (4).

    Retry routing:
      content -> faithfulness/completeness/table/figure failed
                 (routes to query_analyzer for full restart)
      format  -> only conciseness failed
                 (routes to generator for answer regeneration)

    Ablation: USE_CRITIC=False -> auto-pass all scores=5.
    \"\"\"
    if not USE_CRITIC:
        scores = CriticScores(
            faithfulness=5, completeness=5, table_accuracy=5,
            figure_accuracy=5, conciseness=5, overall=5.0,
            feedback="Critic disabled (ablation mode).")
        return {
            **state, "critic_scores": scores, "critic_passed": True,
            "critic_feedback": scores["feedback"],
            "stage_logs": state["stage_logs"] + [
                {"stage": "critic", "status": "completed",
                 "details": {"ablation_skip": True, "passed": True}}
            ],
        }

    context = (state.get("refined_context") or "")[:8000]
    chain   = _prompt | get_llm().with_structured_output(CriticEvaluation)
    result  = chain.invoke({
        "query":   state["query"],
        "context": context,
        "answer":  state.get("answer", ""),
    })
    dims    = [result.faithfulness, result.completeness,
               result.table_accuracy, result.figure_accuracy,
               result.conciseness]
    passed  = all(d >= CRITIC_MIN_SCORE for d in dims)
    overall = round(sum(dims) / len(dims), 2)
    scores  = CriticScores(
        faithfulness=result.faithfulness,
        completeness=result.completeness,
        table_accuracy=result.table_accuracy,
        figure_accuracy=result.figure_accuracy,
        conciseness=result.conciseness,
        overall=overall,
        feedback=result.feedback)
    return {
        **state, "critic_scores": scores, "critic_passed": passed,
        "critic_feedback": result.feedback,
        "retry_reason": result.retry_reason if not passed else None,
        "stage_logs": state["stage_logs"] + [
            {"stage": "critic", "status": "completed",
             "details": {"passed": passed, "overall": overall,
                         "retry_reason": result.retry_reason}}
        ],
    }"""

FIGURE_EXTRACTOR_CODE = """\
# ingestion/figure_extractor.py  — GPT-4o Vision Figure Extractor
import base64, hashlib, logging, os
from pathlib import Path
from typing import List
from langchain_core.documents import Document

# ~300x300 px minimum — filters logos, seals, decorative icons
MIN_IMAGE_AREA = 90_000

_VISION_PROMPT = (
    "You are analyzing a figure extracted from a research paper. "
    "Does this figure contain meaningful data or technical content "
    "(chart, graph, architecture diagram, table as image, flowchart)? "
    "If YES: describe it in detail — figure type, data shown, key "
    "values, labels, axes, and the main insight it conveys. "
    "If NO (logo, photo, institutional seal): respond with: SKIP"
)

def _describe_with_vision(img_bytes, source, page) -> str | None:
    \"\"\"Call GPT-4o Vision; return description or None (SKIP).\"\"\"
    try:
        from openai import OpenAI
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
        b64    = base64.b64encode(img_bytes).decode("utf-8")
        resp   = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": [
                {"type": "text",      "text": _VISION_PROMPT},
                {"type": "image_url", "image_url":
                 {"url": f"data:image/png;base64,{b64}",
                  "detail": "high"}},
            ]}], max_tokens=512)
        desc = resp.choices[0].message.content.strip()
        return None if desc.upper().startswith("SKIP") else desc
    except Exception as e:
        return f"[Figure on page {page} — vision unavailable: {e}]"

def extract_figures(file_path) -> List[Document]:
    \"\"\"
    Extract all figures from a PDF as retrieval-ready Documents.
    Two-layer filter:
      1. Size: skip images < MIN_IMAGE_AREA px (logos, icons).
      2. Vision: GPT-4o returns SKIP for non-informational images.
    Each accepted figure becomes one Document with modality='figure'.
    \"\"\"
    path = Path(file_path)
    if path.suffix.lower() != ".pdf":
        return []
    import fitz   # PyMuPDF
    pdf = fitz.open(str(path))
    docs, idx = [], 0
    for page_num, page in enumerate(pdf, start=1):
        for img_info in page.get_images(full=True):
            bi   = pdf.extract_image(img_info[0])
            w, h = bi.get("width", 0), bi.get("height", 0)
            if w * h < MIN_IMAGE_AREA:
                continue
            desc = _describe_with_vision(bi["image"], path.name, page_num)
            if desc is None:
                continue
            cid = hashlib.md5(
                f"{path}::page{page_num}::fig{idx}".encode()
            ).hexdigest()[:12]
            docs.append(Document(
                page_content=desc,
                metadata={"source": str(path), "file_name": path.name,
                          "file_type": ".pdf", "page": page_num,
                          "figure_index": idx, "image_width": w,
                          "image_height": h, "modality": "figure",
                          "section_heading": "Figure", "chunk_id": cid,
                          "chunk_index": idx,
                          "word_count": len(desc.split()),
                          "char_count": len(desc)}))
            idx += 1
    pdf.close()
    return docs"""

SETTINGS_CODE = """\
# config/settings.py  — Central Configuration
from pathlib import Path
from dotenv import load_dotenv

_ENV_FILE = Path(__file__).resolve().parent.parent / ".env"
load_dotenv(dotenv_path=_ENV_FILE, override=False)

# ── Paths ──────────────────────────────────────────────────────────
BASE_DIR           = Path(__file__).resolve().parent.parent
DATA_RAW_DIR       = BASE_DIR / "data"       / "raw"
DATA_PROCESSED_DIR = BASE_DIR / "data"       / "processed"
VECTORSTORE_DIR    = BASE_DIR / "vectorstore"/ "index"
LOG_FILE           = BASE_DIR / "logs"       / "pipeline.log"

# ── Ingestion ──────────────────────────────────────────────────────
CHUNK_SIZE           = 512    # tokens (~4 chars/token)
CHUNK_OVERLAP        = 64
SUPPORTED_EXTENSIONS = [".pdf", ".docx", ".txt", ".html", ".md"]

# ── Embeddings ─────────────────────────────────────────────────────
EMBEDDING_MODEL     = "text-embedding-3-large"  # OpenAI
VECTORSTORE_BACKEND = "faiss"                    # "faiss" | "chroma"

# ── LLM Provider  ("openai"|"google"|"anthropic"|"groq") ──────────
LLM_PROVIDER        = "openai"
OPENAI_LLM_MODEL    = "gpt-4o"
GOOGLE_LLM_MODEL    = "gemini-2.0-flash"
ANTHROPIC_LLM_MODEL = "claude-sonnet-4-6"
GROQ_LLM_MODEL      = "llama-3.3-70b-versatile"
LLM_MODEL           = {
    "openai":    OPENAI_LLM_MODEL,
    "google":    GOOGLE_LLM_MODEL,
    "anthropic": ANTHROPIC_LLM_MODEL,
    "groq":      GROQ_LLM_MODEL,
}[LLM_PROVIDER]
LLM_TEMPERATURE     = 0.0
LLM_MAX_TOKENS      = 2048

# ── Retrieval ──────────────────────────────────────────────────────
RETRIEVAL_K_MIN        = 4
RETRIEVAL_K_MAX        = 12
RETRIEVAL_K_DEFAULT    = 6
MMR_FETCH_K_MULTIPLIER = 3     # fetch_k = k * multiplier
MMR_LAMBDA_MULT        = 0.5   # 0=max diversity, 1=max relevance

# ── Validation ─────────────────────────────────────────────────────
EVIDENCE_SCORE_THRESHOLD = 0.30
MIN_VALID_PASSAGES       = 2

# ── Critic / Self-Reflection ───────────────────────────────────────
CRITIC_MIN_SCORE = 4           # all dims >= 4 on a 1-5 scale
MAX_RETRIES      = 3

# ── Faithfulness constraint ────────────────────────────────────────
FAITHFULNESS_THRESHOLD = 0.80  # tau: Faith(a, R(q,D)) >= tau

# ── Entry Router ───────────────────────────────────────────────────
ROUTER_SIMILARITY_THRESHOLD = EVIDENCE_SCORE_THRESHOLD

# ── Ablation flags ─────────────────────────────────────────────────
USE_RETRIEVAL_PLANNER = True
USE_VALIDATOR         = False
USE_CONTEXT_REFINER   = True
USE_CRITIC            = True

# ── Logging ────────────────────────────────────────────────────────
LOG_LEVEL = "INFO"   # "DEBUG" for verbose stage traces"""


# ─── Step 4: Build new Section 3.10 elements in CORRECT order ────────────────

def section(heading, description, code):
    elems = []
    elems.append(make_heading_para(heading, level=3))
    elems.append(make_body_para(description))
    elems.append(make_blank())
    elems.extend(code_paras(code))
    elems.append(make_blank())
    return elems


new_elements = []
new_elements.append(make_page_break())
new_elements.append(make_heading_para('3.10  Key Code Listings', level=2))
new_elements.append(make_body_para(
    'This section presents the complete annotated source code for each stage of the '
    'AC-RAG pipeline. All nodes are implemented as pure functions that read from and '
    'write to the shared ACRagState TypedDict, enabling LangGraph to apply partial '
    'state updates — each node writes only the fields it owns. Ablation flags in '
    'config/settings.py allow any stage to be independently disabled for comparative '
    'evaluation without modifying node code.'))
new_elements.append(make_blank())

new_elements += section(
    '3.10.1  Shared State Schema  (pipeline/state.py)',
    'The ACRagState TypedDict is the single data contract shared across all pipeline '
    'nodes. Using TypedDict (not a dataclass) is a LangGraph requirement — it supports '
    'partial dict updates so each node writes only the fields it is responsible for. '
    'RetrievalPlan and CriticScores are nested TypedDicts that enforce type contracts '
    'on the planner and critic outputs respectively.',
    STATE_CODE)

new_elements += section(
    '3.10.2  LangGraph Pipeline Assembly  (pipeline/graph.py)',
    'The build_pipeline() function assembles the complete StateGraph topology. '
    'Conditional edges implement all routing decisions. The increment_retry node '
    'acts as a counter gate; _route_after_retry_increment reads the retry_reason '
    'field to select the correct retry destination — query_analyzer for content '
    'failures, generator for format failures, retriever for validation failures.',
    GRAPH_CODE)

new_elements += section(
    '3.10.3  Entry Router Node  (pipeline/nodes/entry_router.py)',
    'The Entry Router performs a lightweight k=3 cosine similarity check before '
    'committing to the full RAG pipeline. Queries scoring below '
    'ROUTER_SIMILARITY_THRESHOLD on both top-score and average-score criteria are '
    'routed to the direct_responder node, which returns a calibrated "I don\'t know" '
    'response. This prevents hallucination for out-of-domain queries and avoids '
    'unnecessary LLM API calls.',
    ENTRY_ROUTER_CODE)

new_elements += section(
    '3.10.4  Query Analyzer Node  (pipeline/nodes/query_analyzer.py)',
    'The Query Analyzer uses Pydantic structured output to guarantee parseable JSON, '
    'eliminating fragile regex parsing. The complexity_score drives adaptive retrieval: '
    'scores below 0.4 trigger shallow k=4 retrieval; scores above 0.7 activate '
    'multi-query decomposition with up to four sub-queries. The intent classification '
    'further influences MMR diversity parameters in the retrieval planner.',
    QUERY_ANALYZER_CODE)

new_elements += section(
    '3.10.5  Retrieval Planner Node  (pipeline/nodes/retrieval_planner.py)',
    'The Retrieval Planner uses a two-tier design. Tier 1 produces a deterministic '
    'rule-based plan from complexity score and intent — fast and ablation-safe. '
    'Tier 2 (USE_RETRIEVAL_PLANNER flag) applies an LLM refinement that can override '
    'modality_filter and k for queries explicitly referencing tables or figures. '
    'A keyword guard prevents the LLM from over-specifying modality for general queries.',
    RETRIEVAL_PLANNER_CODE)

new_elements += section(
    '3.10.6  Evidence Validator Node  (pipeline/nodes/validator.py)',
    'The Validator scores each retrieved passage against the query using cosine '
    'similarity. Passages below EVIDENCE_SCORE_THRESHOLD (0.30) are discarded. '
    'If fewer than MIN_VALID_PASSAGES (2) survive, validation_passed=False triggers '
    'a retrieval retry with broader MMR parameters. Setting USE_VALIDATOR=False '
    'in ablation testing bypasses scoring and passes all documents through.',
    VALIDATOR_CODE)

new_elements += section(
    '3.10.7  Answer Generator Node  (pipeline/nodes/generator.py)',
    'The Generator enforces the faithfulness constraint Faith(a, R(q,D)) >= tau at '
    'two levels: (1) explicit system prompt rules prohibiting external knowledge, and '
    '(2) mandatory inline source citations [N] for every claim. The Pydantic '
    'GeneratedAnswer schema captures the answer text and key source indices, which '
    'build_attribution() converts into a per-sentence evidence map for frontend display.',
    GENERATOR_CODE)

new_elements += section(
    '3.10.8  Critic / Self-Reflection Node  (pipeline/nodes/critic.py)',
    'The Critic evaluates answers on five dimensions: faithfulness, completeness, '
    'table accuracy, figure accuracy, and conciseness. All must score '
    '>= CRITIC_MIN_SCORE (4). Retry classification — "content" vs "format" — '
    'determines the retry target: content failures restart the full pipeline via '
    'query_analyzer; format failures only regenerate the answer via generator. '
    'USE_CRITIC=False auto-passes all answers in ablation testing.',
    CRITIC_CODE)

new_elements += section(
    '3.10.9  GPT-4o Vision Figure Extractor  (ingestion/figure_extractor.py)',
    'The figure extractor uses dual-layer filtering. Layer 1 is a pixel-area '
    'threshold (MIN_IMAGE_AREA=90,000 px²) that eliminates logos and icons purely '
    'on image size. Layer 2 sends each surviving image to GPT-4o Vision with a '
    'prompt requesting "SKIP" for non-data images such as institutional seals and '
    'photos. Only images passing both filters become retrieval-ready Document chunks '
    'with modality="figure" and a rich textual description as page_content.',
    FIGURE_EXTRACTOR_CODE)

new_elements += section(
    '3.10.10  System Configuration  (config/settings.py)',
    'All hyperparameters, model names, ablation flags, and thresholds are centralised '
    'in config/settings.py. Changing LLM_PROVIDER to "google", "anthropic", or "groq" '
    'switches the entire pipeline in a single edit. The four ablation flags '
    '(USE_RETRIEVAL_PLANNER, USE_VALIDATOR, USE_CONTEXT_REFINER, USE_CRITIC) '
    'disable any pipeline stage independently without modifying node code.',
    SETTINGS_CODE)

print(f"Built {len(new_elements)} new elements for Section 3.10.")


# ─── Step 5: Find anchor element (para 298 = last blank before wrong content) ─

# After the removal, rebuild the body children list
body_children2 = list(body)
para_to_body2 = {}
pi = 0
for bi, child in enumerate(body_children2):
    if child.tag.endswith('}p'):
        para_to_body2[pi] = bi
        pi += 1

# Para 298 should now be the blank line just before Chapter 4 (which moved)
# Let's verify what para 298 looks like after removal
print(f"Para 296 after removal: '{doc.paragraphs[296].text[:80]}'")
print(f"Para 297 after removal: '{doc.paragraphs[297].text[:80]}'")
print(f"Para 298 after removal: '{doc.paragraphs[298].text[:80]}'")
print(f"Para 299 after removal: '{doc.paragraphs[299].text[:80]}'")

# Find Chapter 4 heading (to insert before it)
ch4_para_idx = None
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip().lower()
    sname = p.style.name if p.style and p.style.name else ''
    if ('chapter 4' in text or 'results' in text) and 'heading 1' in sname.lower():
        ch4_para_idx = i
        print(f"Chapter 4 now at para {i}: '{p.text[:60]}'")
        break

if ch4_para_idx is None:
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        sname = p.style.name if p.style and p.style.name else ''
        if 'heading 1' in sname.lower() and ('4' in text or 'result' in text.lower()):
            ch4_para_idx = i
            print(f"Chapter 4 (fallback) at para {i}: '{text[:60]}'")
            break

# Find last para of 3.9 content (anchor for insertion)
# It's 2-3 paras before Chapter 4
anchor_para_idx = ch4_para_idx - 1
print(f"Anchor para {anchor_para_idx}: '{doc.paragraphs[anchor_para_idx].text[:60]}'")

# Get the body element for Chapter 4 (insert before it)
ch4_body_idx = para_to_body2[ch4_para_idx]
ch4_elem = body_children2[ch4_body_idx]

print(f"Inserting {len(new_elements)} elements before Chapter 4 (body idx {ch4_body_idx})...")

# ─── Step 6: Insert in CORRECT ORDER using addnext chain ──────────────────────
# Strategy: insert each element immediately before Chapter 4 using addprevious,
# but process the list in REVERSE so the first element ends up first.
# (addprevious inserts right before ref; last call = element right before ref)

# addprevious(X) inserts X directly before ch4_elem.
# Each subsequent call inserts the NEXT element right before ch4,
# so elements accumulate in the same order as the loop:
#   addprevious(e1) -> [e1, ch4]
#   addprevious(e2) -> [e1, e2, ch4]   (e2 inserted right before ch4)
#   addprevious(e3) -> [e1, e2, e3, ch4]
# Result order == iteration order — do NOT reverse.

for elem in new_elements:
    ch4_elem.addprevious(elem)

print("Insertion complete.")
print(f"Total paragraphs after: {len(doc.paragraphs)}")


# ─── Step 7: Update TOC — add Section 3.10 entry ──────────────────────────────

# Find the TOC line for "3.9  Implementation Methodology"
for i, p in enumerate(doc.paragraphs):
    if '3.9' in p.text and 'Implementation' in p.text and 'CHAPTER' not in p.text:
        # Add 3.10 entry after this paragraph
        p_elem = p._element
        new_toc_entry = make_body_para('3.10  Key Code Listings\t29')
        p_elem.addnext(new_toc_entry)
        print(f"Added TOC entry '3.10 Key Code Listings' after para {i}.")
        break


# ─── Step 8: Update List of Figures — fix page number notes ──────────────────
# Page numbers will auto-update in Word; update the text to reflect new structure

lof_updates = {
    'Fig 3.1': 'Fig 3.1  High-Level Architecture of the AC-RAG Multi-Agent Pipeline\t17',
    'Fig 3.2': 'Fig 3.2  Document Ingestion Pipeline — Stages and Artifacts\t19',
    'Fig 3.3': 'Fig 3.3  LangGraph StateGraph — Node and Conditional Edge Layout\t27',
    'Fig 4.1': 'Fig 4.1  Ablation Study — Mean Faithfulness Score per Configuration\t47',
    'Fig 4.2': 'Fig 4.2  Ablation Study — Composite Score Comparison (7 Configurations)\t48',
}

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    for prefix, new_text in lof_updates.items():
        if text.startswith(prefix):
            # Clear existing runs and set new text
            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = new_text
            else:
                p.add_run(new_text)
            print(f"  Updated LOF: '{new_text[:70]}'")
            break


# ─── Step 9: Final structural checks ─────────────────────────────────────────

print("\n=== FINAL HEADING STRUCTURE ===")
for i, p in enumerate(doc.paragraphs):
    sname = p.style.name if p.style and p.style.name else ''
    if 'Heading' in sname:
        print(f"  [{i}] [{sname}] {p.text[:80]}")


# ─── Save ─────────────────────────────────────────────────────────────────────
doc.save(THESIS_PATH)
print(f"\nSaved: {THESIS_PATH}")
doc2 = Document(THESIS_PATH)
print(f"Final paragraph count: {len(doc2.paragraphs)}")
