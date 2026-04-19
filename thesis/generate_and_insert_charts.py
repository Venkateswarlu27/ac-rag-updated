"""
generate_and_insert_charts.py
Creates Fig 4.1 and Fig 4.2 as PNG charts and inserts them into the thesis
before their respective captions (paragraphs 952 and 953).
"""

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

THESIS_PATH = "AC_RAG_Thesis_Submission_Ready.docx"

# ─── Ablation data ────────────────────────────────────────────────────────────

configs = [
    "Full AC-RAG",
    "w/o Self-\nReflection",
    "w/o Evidence\nValidation",
    "w/o Adaptive\nRetrieval",
    "w/o Multi-Modal\nRetrieval",
    "w/o Context\nRefinement",
    "w/o Router\n(always RAG)",
]

AF = [0.88, 0.73, 0.75, 0.80, 0.79, 0.82, 0.88]
HR = [0.85, 0.70, 0.72, 0.77, 0.76, 0.78, 0.85]
CR = [0.86, 0.86, 0.72, 0.79, 0.71, 0.86, 0.86]
RA = [0.87, 0.74, 0.76, 0.81, 0.78, 0.83, 0.86]

# ─── Fig 4.1: Faithfulness bar chart ─────────────────────────────────────────

fig1, ax1 = plt.subplots(figsize=(11, 5.5))

colors = ['#2E86AB' if v < 0.88 else '#E84855' for v in AF]
bars = ax1.bar(range(len(configs)), AF, color=colors, edgecolor='white',
               linewidth=0.8, width=0.6)

# Value labels on top of bars
for bar, val in zip(bars, AF):
    ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.005,
             f'{val:.2f}', ha='center', va='bottom', fontsize=11, fontweight='bold',
             color='#333333')

ax1.set_xticks(range(len(configs)))
ax1.set_xticklabels(configs, fontsize=9.5)
ax1.set_ylim(0.60, 0.96)
ax1.set_ylabel('Mean Answer Faithfulness Score', fontsize=12)
ax1.set_title('Figure 4.1: Ablation Study — Mean Faithfulness Score per Configuration',
              fontsize=12, fontweight='bold', pad=14)
ax1.axhline(y=0.88, color='#E84855', linestyle='--', linewidth=1.2, alpha=0.6,
            label='Full AC-RAG (0.88)')
ax1.set_xlabel('Pipeline Configuration', fontsize=11)
ax1.yaxis.grid(True, linestyle='--', alpha=0.4)
ax1.set_axisbelow(True)
ax1.spines['top'].set_visible(False)
ax1.spines['right'].set_visible(False)

legend_patch = mpatches.Patch(color='#E84855', label='Full AC-RAG (0.88)')
ablation_patch = mpatches.Patch(color='#2E86AB', label='Ablation variants')
ax1.legend(handles=[legend_patch, ablation_patch], fontsize=10, loc='lower right')

plt.tight_layout()
buf1 = io.BytesIO()
fig1.savefig(buf1, format='png', dpi=150, bbox_inches='tight')
buf1.seek(0)
plt.close(fig1)
print("Fig 4.1 generated.")

# ─── Fig 4.2: Composite grouped bar chart ────────────────────────────────────

fig2, ax2 = plt.subplots(figsize=(13, 6))

x = np.arange(len(configs))
width = 0.20

metric_colors = ['#E84855', '#2E86AB', '#06A77D', '#F4A261']

b1 = ax2.bar(x - 1.5*width, AF, width, label='Answer Faithfulness (AF)',
             color=metric_colors[0], edgecolor='white', linewidth=0.5)
b2 = ax2.bar(x - 0.5*width, HR, width, label='Hallucination Reduction (HR)',
             color=metric_colors[1], edgecolor='white', linewidth=0.5)
b3 = ax2.bar(x + 0.5*width, CR, width, label='Context Relevance (CR)',
             color=metric_colors[2], edgecolor='white', linewidth=0.5)
b4 = ax2.bar(x + 1.5*width, RA, width, label='Response Accuracy (RA)',
             color=metric_colors[3], edgecolor='white', linewidth=0.5)

# Value labels
for bars_group in [b1, b2, b3, b4]:
    for bar in bars_group:
        h = bar.get_height()
        ax2.text(bar.get_x() + bar.get_width()/2, h + 0.003,
                 f'{h:.2f}', ha='center', va='bottom', fontsize=7.5,
                 fontweight='bold', color='#333333', rotation=0)

ax2.set_xticks(x)
ax2.set_xticklabels(configs, fontsize=8.5)
ax2.set_ylim(0.58, 0.98)
ax2.set_ylabel('Mean Score', fontsize=12)
ax2.set_title('Figure 4.2: Ablation Study — Composite Score Comparison (7 Configurations)',
              fontsize=12, fontweight='bold', pad=14)
ax2.set_xlabel('Pipeline Configuration', fontsize=11)
ax2.yaxis.grid(True, linestyle='--', alpha=0.4)
ax2.set_axisbelow(True)
ax2.spines['top'].set_visible(False)
ax2.spines['right'].set_visible(False)
ax2.legend(loc='lower right', fontsize=9.5, ncol=2, framealpha=0.9)

plt.tight_layout()
buf2 = io.BytesIO()
fig2.savefig(buf2, format='png', dpi=150, bbox_inches='tight')
buf2.seek(0)
plt.close(fig2)
print("Fig 4.2 generated.")

# ─── Insert charts into thesis ────────────────────────────────────────────────

doc = Document(THESIS_PATH)

# Find paragraphs with Fig 4.1 and Fig 4.2 captions
fig41_idx = None
fig42_idx = None
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if 'Fig 4.1' in text and 'Faithfulness' in text:
        fig41_idx = i
    if 'Fig 4.2' in text and 'Composite' in text:
        fig42_idx = i

print(f"Fig 4.1 caption at paragraph {fig41_idx}")
print(f"Fig 4.2 caption at paragraph {fig42_idx}")


def insert_image_before_para(doc, para_idx, image_buf, width_inches=5.8):
    """Insert an image paragraph immediately before the paragraph at para_idx."""
    body = doc.element.body
    body_children = list(body)

    # Map paragraph index -> body child index
    pi = 0
    target_body_idx = None
    for bi, child in enumerate(body_children):
        if child.tag.endswith('}p'):
            if pi == para_idx:
                target_body_idx = bi
                break
            pi += 1

    if target_body_idx is None:
        print(f"  WARNING: Could not find body element for para {para_idx}")
        return

    ref_elem = body_children[target_body_idx]

    # Build image paragraph using python-docx internals
    from docx.shared import Inches
    import docx.oxml.ns as ns

    # Create a temporary paragraph via doc to get the run/image XML
    temp_para = doc.add_paragraph()
    run = temp_para.add_run()
    run.add_picture(image_buf, width=Inches(width_inches))

    # Move the temp para element before the reference element
    temp_elem = temp_para._element
    temp_elem.getparent().remove(temp_elem)
    ref_elem.addprevious(temp_elem)

    # Add a small space paragraph between image and caption
    space_para = doc.add_paragraph()
    space_elem = space_para._element
    space_elem.getparent().remove(space_elem)
    ref_elem.addprevious(space_elem)

    print(f"  Image inserted before paragraph {para_idx}")


if fig41_idx is not None:
    buf1.seek(0)
    insert_image_before_para(doc, fig41_idx, buf1, width_inches=5.8)

if fig42_idx is not None:
    buf2.seek(0)
    insert_image_before_para(doc, fig42_idx, buf2, width_inches=6.2)

doc.save(THESIS_PATH)
print(f"Saved: {THESIS_PATH}")
print(f"Total paragraphs: {len(Document(THESIS_PATH).paragraphs)}")
